import os
import sys
import time
import json
import logging
import argparse
from pathlib import Path
import subprocess
import whisper
import yt_dlp
from tqdm import tqdm
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re
import base64

def sanitize_filename(filename):
    """Remove caracteres inválidos de nomes de arquivos."""
    invalid_chars = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    if len(filename) > 200:
        filename = filename[:200]
    return filename

class YouTubeAnalyzerWeb:
    def __init__(self, temp_dir: str = "temp"):
        self.temp_dir = Path(temp_dir)
        self.ffmpeg_path = Path("ffmpeg.exe")
        self.ffprobe_path = Path("ffprobe.exe")
        
        # Configurar logging para modo web
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)
        
        # Carregar modelo Whisper
        self.logger.info("Carregando modelo Whisper...")
        self.model = whisper.load_model("base")
        
        # Inicializar o ambiente
        self._setup_environment()
        
        # Configurar opções do yt-dlp
        self.ydl_opts = {
            'format': 'worstaudio',
            'outtmpl': str(self.temp_dir / '%(title)s.%(ext)s'),
            'ffmpeg_location': str(self.ffmpeg_path.parent),
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'wav',
                'preferredquality': '96',
            }],
            'quiet': True,
            'no_warnings': True
        }
    
    def _setup_environment(self) -> None:
        """Configura o ambiente necessário para o analisador."""
        try:
            self.temp_dir.mkdir(parents=True, exist_ok=True)
            self._check_ffmpeg()
        except Exception as e:
            self.logger.error(f"Erro na configuração do ambiente: {str(e)}")
            raise
    
    def _check_ffmpeg(self) -> None:
        """Verifica se os executáveis do FFmpeg estão disponíveis."""
        if not self.ffmpeg_path.exists():
            raise FileNotFoundError(f"FFmpeg não encontrado em {self.ffmpeg_path}")
        if not self.ffprobe_path.exists():
            raise FileNotFoundError(f"FFprobe não encontrado em {self.ffprobe_path}")
    
    def download_video(self, url: str) -> tuple[bool, str]:
        """Baixa o vídeo do YouTube e retorna o caminho para o arquivo de áudio."""
        try:
            with yt_dlp.YoutubeDL(self.ydl_opts) as ydl:
                info = ydl.extract_info(url, download=False)
                title = info.get('title', 'video')
                
                self.logger.info(f"Baixando: {title}")
                ydl.download([url])
                
                # Encontrar o arquivo de áudio baixado
                audio_files = list(self.temp_dir.glob("*.wav"))
                if audio_files:
                    return True, str(audio_files[0])
                else:
                    return False, "Arquivo de áudio não encontrado"
                    
        except Exception as e:
            self.logger.error(f"Erro no download: {str(e)}")
            return False, str(e)
    
    def transcribe_audio(self, audio_path: str) -> tuple[bool, str]:
        """Transcreve o áudio usando o modelo Whisper."""
        try:
            self.logger.info("Iniciando transcrição...")
            result = self.model.transcribe(audio_path, language="pt")
            return True, result["text"]
        except Exception as e:
            self.logger.error(f"Erro na transcrição: {str(e)}")
            return False, str(e)
    
    def generate_summary(self, transcription: str) -> str:
        """Gera um resumo básico da transcrição."""
        # Resumo simples - primeiras 500 palavras
        words = transcription.split()
        if len(words) > 500:
            summary = " ".join(words[:500]) + "..."
        else:
            summary = transcription
        
        return f"Resumo do conteúdo:\n\n{summary}"
    
    def create_document(self, title: str, transcription: str, summary: str, request_id: str) -> str:
        """Cria um documento DOCX formatado."""
        try:
            doc = Document()
            
            # Título
            title_paragraph = doc.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data
            date_paragraph = doc.add_paragraph(f"Data de análise: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_page_break()
            
            # Resumo
            doc.add_heading('Resumo', level=1)
            doc.add_paragraph(summary)
            
            doc.add_page_break()
            
            # Transcrição completa
            doc.add_heading('Transcrição Completa', level=1)
            doc.add_paragraph(transcription)
            
            # Salvar documento
            safe_title = sanitize_filename(title)
            doc_path = self.temp_dir / f"{safe_title}_{request_id}.docx"
            doc.save(str(doc_path))
            
            return str(doc_path)
            
        except Exception as e:
            self.logger.error(f"Erro ao criar documento: {str(e)}")
            raise
    
    def cleanup(self, files_to_remove: list) -> None:
        """Remove arquivos temporários."""
        for file_path in files_to_remove:
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
            except Exception as e:
                self.logger.warning(f"Erro ao remover {file_path}: {str(e)}")
    
    def process_video_web(self, url: str, request_id: str) -> dict:
        """Processa um vídeo para a interface web."""
        files_to_cleanup = []
        
        try:
            # Download do vídeo
            success, audio_path = self.download_video(url)
            if not success:
                return {"error": f"Erro no download: {audio_path}"}
            
            files_to_cleanup.append(audio_path)
            
            # Obter título do vídeo
            with yt_dlp.YoutubeDL({'quiet': True}) as ydl:
                info = ydl.extract_info(url, download=False)
                title = info.get('title', 'Vídeo do YouTube')
            
            # Transcrição
            success, transcription = self.transcribe_audio(audio_path)
            if not success:
                return {"error": f"Erro na transcrição: {transcription}"}
            
            # Gerar resumo
            summary = self.generate_summary(transcription)
            
            # Criar documento
            doc_path = self.create_document(title, transcription, summary, request_id)
            
            # Converter documento para base64 para download
            with open(doc_path, 'rb') as f:
                doc_base64 = base64.b64encode(f.read()).decode('utf-8')
            
            files_to_cleanup.append(doc_path)
            
            result = {
                "title": title,
                "transcription": transcription,
                "summary": summary,
                "downloadUrl": f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{doc_base64}"
            }
            
            return result
            
        except Exception as e:
            self.logger.error(f"Erro no processamento: {str(e)}")
            return {"error": str(e)}
        
        finally:
            # Limpeza
            self.cleanup(files_to_cleanup)

def main():
    parser = argparse.ArgumentParser(description='Analisador de Vídeos do YouTube - Modo Web')
    parser.add_argument('--web-mode', action='store_true', help='Modo web')
    parser.add_argument('--input', required=True, help='Arquivo com URL de entrada')
    
    args = parser.parse_args()
    
    if not args.web_mode:
        print("Este script deve ser executado em modo web")
        sys.exit(1)
    
    try:
        # Ler URL do arquivo
        with open(args.input, 'r') as f:
            url = f.read().strip()
        
        # Extrair request_id do nome do arquivo
        request_id = Path(args.input).stem.split('_')[-1]
        
        # Processar vídeo
        analyzer = YouTubeAnalyzerWeb()
        result = analyzer.process_video_web(url, request_id)
        
        # Salvar resultado
        result_file = Path('temp') / f'result_{request_id}.json'
        with open(result_file, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        
        print("Processamento concluído com sucesso")
        
    except Exception as e:
        print(f"Erro: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()