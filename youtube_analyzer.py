import os
import time
import logging
from pathlib import Path
import subprocess
import whisper
import yt_dlp
from tqdm import tqdm
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re  # Adicione esta importação para usar expressões regulares

def sanitize_filename(filename):
    """Remove caracteres inválidos de nomes de arquivos."""
    # Substituir caracteres inválidos por underscores
    invalid_chars = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    # Limitar o tamanho do nome do arquivo
    if len(filename) > 200:
        filename = filename[:200]
    return filename

class YouTubeAnalyzer:
    def __init__(self, temp_dir: str = "temp"):
        self.temp_dir = Path(temp_dir)
        self.ffmpeg_path = Path("ffmpeg.exe")
        self.ffprobe_path = Path("ffprobe.exe")
        
        # Configurar logging
        logging.basicConfig(
            level=logging.DEBUG,  # Alterado para DEBUG para mais detalhes
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)
        
        self.logger.info("Inicializando YouTubeAnalyzer")
        self.logger.debug(f"Diretório temporário: {self.temp_dir}")
        self.logger.debug(f"FFmpeg path: {self.ffmpeg_path}")
        
        # Carregar modelo Whisper
        self.logger.info("Carregando modelo Whisper...")
        start_time = time.time()
        self.model = whisper.load_model("base")
        load_time = time.time() - start_time
        self.logger.info(f"Modelo Whisper carregado em {load_time:.2f} segundos")
        
        # Inicializar o ambiente
        self._setup_environment()
        
        # Configurar opções do yt-dlp
        self.logger.debug("Configurando opções do yt-dlp")
        self.ydl_opts = {
            'format': 'worstaudio',  # Baixar apenas áudio na menor qualidade
            'outtmpl': str(self.temp_dir / '%(title)s.%(ext)s'),
            'ffmpeg_location': str(self.ffmpeg_path.parent),
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'wav',
                'preferredquality': '96',
            }],
            'quiet': True,
            'no_warnings': True
        }
        self.logger.debug(f"Opções do yt-dlp configuradas: {self.ydl_opts}")
    
    def _setup_environment(self) -> None:
        """Configura o ambiente necessário para o analisador."""
        try:
            self.logger.debug("Iniciando configuração do ambiente")
            # Criar diretório temporário se não existir
            self.temp_dir.mkdir(parents=True, exist_ok=True)
            self.logger.debug(f"Diretório temporário verificado: {self.temp_dir}")
            
            # Verificar dependências do FFmpeg
            self._check_ffmpeg()
            
            self.logger.info("Ambiente configurado com sucesso")
        except Exception as e:
            self.logger.error(f"Erro na configuração do ambiente: {str(e)}")
            raise
    
    def _check_ffmpeg(self) -> None:
        """Verifica se os executáveis do FFmpeg estão disponíveis."""
        self.logger.debug("Verificando executáveis do FFmpeg")
        if not self.ffmpeg_path.exists():
            self.logger.error(f"FFmpeg não encontrado em {self.ffmpeg_path}")
            raise FileNotFoundError(f"FFmpeg não encontrado em {self.ffmpeg_path}")
        if not self.ffprobe_path.exists():
            self.logger.error(f"FFprobe não encontrado em {self.ffprobe_path}")
            raise FileNotFoundError(f"FFprobe não encontrado em {self.ffprobe_path}")
        self.logger.debug("Executáveis do FFmpeg verificados com sucesso")
    
    def download_video(self, url: str) -> tuple[bool, str]:
        """Baixa o vídeo do YouTube e retorna o caminho para o arquivo de áudio."""
        try:
            self.logger.info(f"Iniciando download do vídeo: {url}")
            start_time = time.time()
            
            with yt_dlp.YoutubeDL(self.ydl_opts) as ydl:
                self.logger.debug("Extraindo informações do vídeo")
                info = ydl.extract_info(url, download=False)
                video_title = info.get('title', 'video')
                self.logger.info(f"Título do vídeo: {video_title}")
                
                # Registrar arquivos WAV existentes antes do download
                existing_wav_files = set(self.temp_dir.glob("*.wav"))
                
                # Baixar vídeo e extrair áudio
                self.logger.info("Baixando vídeo e extraindo áudio...")
                ydl.download([url])
                
                # Encontrar novos arquivos WAV após o download
                current_wav_files = set(self.temp_dir.glob("*.wav"))
                new_wav_files = current_wav_files - existing_wav_files
                
                if new_wav_files:
                    # Usar o arquivo WAV recém-baixado
                    audio_path = list(new_wav_files)[0]
                    download_time = time.time() - start_time
                    file_size = audio_path.stat().st_size / (1024 * 1024)  # MB
                    self.logger.info(f"Download concluído em {download_time:.2f} segundos")
                    self.logger.debug(f"Arquivo salvo em: {audio_path}")
                    self.logger.debug(f"Tamanho do arquivo: {file_size:.2f} MB")
                    return True, str(audio_path)
                else:
                    # Tentar encontrar o arquivo pelo título do vídeo
                    expected_file = self.temp_dir / f"{video_title}.wav"
                    if expected_file.exists():
                        download_time = time.time() - start_time
                        file_size = expected_file.stat().st_size / (1024 * 1024)  # MB
                        self.logger.info(f"Download concluído em {download_time:.2f} segundos")
                        self.logger.debug(f"Arquivo salvo em: {expected_file}")
                        self.logger.debug(f"Tamanho do arquivo: {file_size:.2f} MB")
                        return True, str(expected_file)
                    else:
                        self.logger.debug(f"Arquivos na pasta temp: {list(self.temp_dir.glob('*'))}")
                        self.logger.error("Nenhum arquivo de áudio WAV novo encontrado após o download")
                        return False, ""
        except Exception as e:
            self.logger.error(f"Erro ao baixar o vídeo: {str(e)}")
            # Adicionar mensagens mais descritivas em português
            if "HTTP Error 404" in str(e):
                self.logger.error("Vídeo não encontrado. Verifique se a URL está correta.")
            elif "Private video" in str(e):
                self.logger.error("Este vídeo é privado e não pode ser acessado.")
            elif "This video is unavailable" in str(e):
                self.logger.error("Este vídeo não está disponível para download.")
            return False, ""
    
    def transcribe_audio(self, audio_path: str) -> tuple[bool, str]:
        """Transcreve o áudio usando o modelo Whisper."""
        try:
            self.logger.info(f"Iniciando transcrição do áudio: {audio_path}")
            start_time = time.time()
            
            audio_file = Path(audio_path)
            if not audio_file.exists():
                self.logger.error(f"Arquivo de áudio não encontrado: {audio_path}")
                return False, ""
            
            # Obter informações do áudio usando ffprobe
            self.logger.debug("Obtendo informações do áudio com ffprobe")
            cmd = [
                str(self.ffprobe_path),
                "-v", "error",
                "-show_entries", "format=duration",
                "-of", "default=noprint_wrappers=1:nokey=1",
                str(audio_file)
            ]
            result = subprocess.run(cmd, capture_output=True, text=True)
            duration = float(result.stdout.strip())
            self.logger.info(f"Duração do áudio: {duration:.2f} segundos")
            
            # Dividir áudio em segmentos para transcrição
            segment_length = 10 * 60  # 10 minutos em segundos
            num_segments = int(duration / segment_length) + 1
            self.logger.info(f"Dividindo áudio em {num_segments} segmentos")
            
            segments = []
            for i in range(num_segments):
                start = i * segment_length
                end = min((i + 1) * segment_length, duration)
                if end - start < 1:  # Ignorar segmentos muito curtos
                    continue
                    
                segment_file = self.temp_dir / f"segment_{i}.wav"
                self.logger.debug(f"Criando segmento {i}: {start:.2f}s - {end:.2f}s")
                
                # Extrair segmento usando ffmpeg
                cmd = [
                    str(self.ffmpeg_path),
                    "-y",  # Sobrescrever arquivo se existir
                    "-i", str(audio_file),
                    "-ss", str(start),
                    "-to", str(end),
                    "-c:a", "pcm_s16le",
                    str(segment_file)
                ]
                subprocess.run(cmd, capture_output=True)
                
                if segment_file.exists():
                    segments.append(str(segment_file))
                    self.logger.debug(f"Segmento {i} criado: {segment_file}")
                else:
                    self.logger.error(f"Falha ao criar segmento {i}")
            
            # Transcrever cada segmento
            self.logger.info(f"Iniciando transcrição de {len(segments)} segmentos")
            transcriptions = []
            
            for i, segment_file in enumerate(tqdm(segments, desc="Transcrevendo segmentos")):
                self.logger.info(f"Transcrevendo segmento {i+1}/{len(segments)}")
                segment_start = time.time()
                
                # Transcrever com Whisper
                # No método transcribe_audio
                # Detectar idioma automaticamente
                detect_result = self.model.transcribe(segment_file, language=None, task="detect_language")
                detected_language = detect_result.get("language", "pt")
                self.logger.info(f"Idioma detectado: {detected_language}")
                
                # Transcrever com o idioma detectado
                result = self.model.transcribe(segment_file, language=detected_language)
                transcription = result["text"]
                
                # No método transcribe_audio, após cada segmento
                # Salvar transcrição intermediária
                intermediate_file = self.temp_dir / f"segment_{i}_transcricao.txt"
                with open(intermediate_file, "w", encoding="utf-8") as f:
                    f.write(transcription)
                self.logger.debug(f"Transcrição intermediária salva em: {intermediate_file}")
                segment_time = time.time() - segment_start
                self.logger.debug(f"Segmento {i+1} transcrito em {segment_time:.2f} segundos")
                self.logger.debug(f"Tamanho da transcrição: {len(transcription)} caracteres")
                
                transcriptions.append(transcription)
            
            # Remover arquivos temporários de segmentos
            self.logger.debug("Removendo arquivos temporários de segmentos")
            for segment_file in segments:
                try:
                    os.remove(segment_file)
                    self.logger.debug(f"Removido: {segment_file}")
                except Exception as e:
                    self.logger.warning(f"Não foi possível remover {segment_file}: {str(e)}")
            
            # Juntar todas as transcrições
            full_transcription = "\n\n".join(transcriptions)
            
            transcription_time = time.time() - start_time
            self.logger.info(f"Transcrição completa em {transcription_time:.2f} segundos")
            self.logger.debug(f"Tamanho total da transcrição: {len(full_transcription)} caracteres")
            
            return True, full_transcription
        except Exception as e:
            self.logger.error(f"Erro ao transcrever o áudio: {str(e)}")
            return False, ""
    
    def save_transcription_to_docx(self, transcription: str, video_title: str) -> tuple[bool, str]:
        """Salva a transcrição em um arquivo Word formatado."""
        try:
            self.logger.info("Iniciando salvamento da transcrição em DOCX")
            start_time = time.time()
            
            # Criar documento Word
            doc = Document()
            
            # Configurar margens do documento
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Adicionar título
            title = doc.add_heading(f"Transcrição: {video_title}", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Adicionar data e hora
            date_time = doc.add_paragraph()
            date_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_time_run = date_time.add_run(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
            date_time_run.font.size = Pt(10)
            date_time_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Adicionar linha separadora
            doc.add_paragraph("_" * 50)
            
            # Adicionar texto da transcrição
            self.logger.debug("Formatando texto da transcrição")
            self.logger.debug(f"Tamanho da transcrição: {len(transcription)} caracteres")
            paragraphs = transcription.split("\n")
            self.logger.debug(f"Número de parágrafos: {len(paragraphs)}")
            for p in paragraphs:
                if p.strip():  # Ignorar linhas vazias
                    doc.add_paragraph(p.strip())
            
            # Salvar documento
            # Sanitizar o título do vídeo para o nome do arquivo
            safe_title = sanitize_filename(video_title)
            self.logger.debug(f"Título sanitizado: {safe_title}")
            
            docx_filename = f"{safe_title}_transcricao.docx"
            docx_path = self.temp_dir / docx_filename
            self.logger.debug(f"Salvando DOCX em: {docx_path}")
            doc.save(str(docx_path))
            
            save_time = time.time() - start_time
            self.logger.info(f"Documento DOCX salvo em {save_time:.2f} segundos")
            self.logger.info(f"Caminho do arquivo: {docx_path}")
            
            if docx_path.exists():
                file_size = docx_path.stat().st_size / 1024  # KB
                self.logger.debug(f"Tamanho do arquivo DOCX: {file_size:.2f} KB")
                return True, str(docx_path)
            else:
                self.logger.error("Arquivo DOCX não foi criado")
                return False, ""
        except Exception as e:
            self.logger.error(f"Erro ao salvar transcrição em DOCX: {str(e)}")
            # Adicionar mais detalhes sobre o erro
            import traceback
            self.logger.error(f"Detalhes do erro: {traceback.format_exc()}")
            return False, ""
    
    def cleanup(self) -> None:
        """Remove arquivos temporários."""
        try:
            self.logger.info("Iniciando limpeza de arquivos temporários")
            # Implementar limpeza conforme necessário
            self.logger.info("Limpeza concluída")
        except Exception as e:
            self.logger.error(f"Erro durante a limpeza: {str(e)}")

# Exemplo de uso
if __name__ == "__main__":
    # Solicitar URL do vídeo ao usuário
    video_url = input("Digite a URL do vídeo do YouTube: ")
    
    # Inicializar analisador
    analyzer = YouTubeAnalyzer()
    
    # Baixar vídeo
    success, audio_path = analyzer.download_video(video_url)
    
    if success:
        # Extrair título do vídeo do caminho do arquivo
        video_title = Path(audio_path).stem
        
        # Transcrever áudio
        success, transcription = analyzer.transcribe_audio(audio_path)
        
        if success:
            # Salvar transcrição em DOCX
            success, docx_path = analyzer.save_transcription_to_docx(transcription, video_title)
            
            if success:
                print(f"\nTranscrição salva em: {docx_path}")
            else:
                print("\nErro ao salvar a transcrição em DOCX")
        else:
            print("\nErro ao transcrever o áudio")
    else:
        print("\nErro ao baixar o vídeo")


def generate_summary(self, transcription: str) -> str:
    """Gera um resumo da transcrição (versão simples)."""
    # Versão simples: pegar as primeiras frases até um limite de caracteres
    sentences = re.split(r'(?<=[.!?]) +', transcription)
    summary = ""
    char_count = 0
    max_chars = 500
    
    for sentence in sentences:
        if char_count + len(sentence) <= max_chars:
            summary += sentence + " "
            char_count += len(sentence) + 1
        else:
            break
            
    return summary.strip() + "..."